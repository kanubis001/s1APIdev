# 가장 기본적인 기능(문서 열기, 저장, 글자 쓰기 등등)
import docx
from docx import Document
from docx.shared import Pt, Cm, Inches
import os
import datetime
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT



class mkReport:
    def __init__(self,startday,endday,names,rootpath):
        self.doc = Document()
        style = self.doc.styles['Normal']
        font = style.font
        font.name = '맑은 고딕'
        font.size = Pt(10)
        title='센티넬원 리포트'
        title2='기간 : ',startday,' ~ ',endday
        if len(names)==2:
            self.scopename=names[0]+">"+names[1]
        elif len(names)==3:
            self.scopename=names[0]+">"+names[1]+">"+names[2]
        title3="scope : "+self.scopename
        self.doc.add_heading(title, level=0)
        self.doc.add_heading(title2, level=1)
        self.doc.add_heading(title3, level=1)
        self.startday=startday
        self.endday=endday
        self.rootpath=rootpath
       
    
    def setValue(self,stype):
        pass

    def mkTable_detail(self, row, col):
        self.table = self.doc.add_table(rows=row, cols=col)
        self.table.style = self.doc.styles['Table Grid']
        self.table.alignment=WD_TABLE_ALIGNMENT.LEFT
        self.table.rows[0].cells[0].text = "agent 명"
        self.table.rows[0].cells[1].text = "IP 주소"
        self.table.rows[0].cells[2].text = "종류"
        self.table.rows[0].cells[3].text = "처리 여부"
        self.table.rows[0].cells[4].text = "위협 구분"
        self.table.rows[0].cells[5].text = "account>site>group"
        
    def mktable_thDetail(self, row, col):
        self.table_thDetail = self.doc.add_table(rows=row, cols=col)
        self.table_thDetail.style=self.doc.styles['Table Grid']
        self.table_thDetail.alignment=WD_TABLE_ALIGNMENT.LEFT
        self.table_thDetail.rows[0].cells[0].width=Cm(4)
        self.table_thDetail.rows[0].cells[1].width=Cm(11.5)
    
    def mktable_getSolved(self, row, col):
        self.table_getSolved = self.doc.add_table(rows=row, cols=col)
        self.table_getSolved.style = self.doc.styles['Table Grid']
        self.table_getSolved.alignment=WD_TABLE_ALIGNMENT.LEFT
        self.table_getSolved.rows[0].cells[0].width=Cm(4)
        self.table_getSolved.rows[0].cells[1].width=Cm(4)
        para=self.table_getSolved.rows[0].cells[0].paragraphs[0]
        para2=self.table_getSolved.rows[0].cells[1].paragraphs[0]
        run=para.add_run('해결 상태')
        run2=para2.add_run("엔드포인트 수량")
        run.bold=True
        run2.bold=True   
        
    def mkTop20Table_detail(self, row, col):
        self.top20table = self.doc.add_table(rows=row, cols=col)
        self.top20table.style = self.doc.styles['Table Grid']
        self.top20table.alignment=WD_TABLE_ALIGNMENT.LEFT
        self.top20table.rows[0].cells[0].width=Cm(1.5)
        self.top20table.rows[0].cells[1].width=Cm(7.5)
        self.top20table.rows[0].cells[2].width=Cm(2.5)
        para=self.top20table.rows[0].cells[0].paragraphs[0]
        para2=self.top20table.rows[0].cells[1].paragraphs[0]
        para3=self.top20table.rows[0].cells[2].paragraphs[0]
        run=para.add_run('No.')
        run2=para2.add_run("Endpoint")
        run3=para3.add_run("탐지건수")
        run.bold=True
        run2.bold=True
        run3.bold=True
        # self.top20table.rows[0].cells[0].text = "No."
        # self.top20table.rows[0].cells[1].text = "Endpoint"
        # self.top20table.rows[0].cells[2].text = "탐지건수"
        
       
    def mktable_getbyos(self, row, col):
        self.table_getbyos = self.doc.add_table(rows=row, cols=col)
        self.table_getbyos.style = self.doc.styles['Table Grid']
        self.table_getbyos.alignment=WD_TABLE_ALIGNMENT.LEFT
        self.table_getbyos.rows[0].cells[0].width=Cm(4)
        self.table_getbyos.rows[0].cells[1].width=Cm(4)
        para=self.table_getbyos.rows[0].cells[0].paragraphs[0]
        para2=self.table_getbyos.rows[0].cells[1].paragraphs[0]
        run=para.add_run('OS 유형')
        run2=para2.add_run("엔드포인트 수량")
        run.bold=True
        run2.bold=True
        # self.table_getbyos.rows[0].cells[0].text = "OS 유형"
        # self.table_getbyos.rows[0].cells[1].text = "엔드포인트 수량"
 
 
    def mktable_getbyver(self, row, col):
        self.table_getbyver = self.doc.add_table(rows=row, cols=col)
        self.table_getbyver.style = self.doc.styles['Table Grid']
        self.table_getbyver.alignment=WD_TABLE_ALIGNMENT.LEFT
        self.table_getbyver.rows[0].cells[0].width=Cm(4)
        self.table_getbyver.rows[0].cells[1].width=Cm(4)
        para=self.table_getbyver.rows[0].cells[0].paragraphs[0]
        para2=self.table_getbyver.rows[0].cells[1].paragraphs[0]
        run=para.add_run('에이전트 버전')
        run2=para2.add_run("엔드포인트 수량")
        run.bold=True
        run2.bold=True
        # self.table_getbyver.rows[0].cells[0].text = "에이전트 버전"
        # self.table_getbyver.rows[0].cells[1].text = "엔드포인트 수량"
        

    def mktable_getbyinf(self, row, col):
        self.table_getbyinf = self.doc.add_table(rows=row, cols=col)
        self.table_getbyinf.style = self.doc.styles['Table Grid']
        self.table_getbyinf.alignment=WD_TABLE_ALIGNMENT.LEFT
        self.table_getbyinf.rows[0].cells[0].width=Cm(4)
        self.table_getbyinf.rows[0].cells[1].width=Cm(4)
        para=self.table_getbyinf.rows[0].cells[0].paragraphs[0]
        para2=self.table_getbyinf.rows[0].cells[1].paragraphs[0]
        run=para.add_run('감염여부')
        run2=para2.add_run("엔드포인트 수량")
        run.bold=True
        run2.bold=True
        # self.table_getbyinf.rows[0].cells[0].text = "감염여부"
        # self.table_getbyinf.rows[0].cells[1].text = "엔드포인트 수량"
        
             
    def mkTblthreatpie(self, row, col):
        self.table_Tblclasspie = self.doc.add_table(rows=row+1, cols=col)
        self.table_Tblclasspie.style = self.doc.styles['Table Grid']
        self.table_Tblclasspie.alignment=WD_TABLE_ALIGNMENT.LEFT
        self.table_Tblclasspie.rows[0].cells[0].width=Cm(6)
        self.table_Tblclasspie.rows[0].cells[1].width=Cm(3)
        para=self.table_Tblclasspie.rows[0].cells[0].paragraphs[0]
        para2=self.table_Tblclasspie.rows[0].cells[1].paragraphs[0]
        run=para.add_run('탐지내역')
        run2=para2.add_run("탐지수")
        run.bold=True
        run2.bold=True
        # self.table_Tblclasspie.rows[0].cells[0].text = "탐지내역"
        # self.table_Tblclasspie.rows[0].cells[1].text = "탐지수"

        
    def mkTblenginechart(self, row, col):
        self.table_Tblengine = self.doc.add_table(rows=row+1, cols=col)
        self.table_Tblengine.style = self.doc.styles['Table Grid']
        self.table_Tblengine.alignment=WD_TABLE_ALIGNMENT.LEFT
        self.table_Tblengine.rows[0].cells[0].width=Cm(6)
        self.table_Tblengine.rows[0].cells[1].width=Cm(3)
        para=self.table_Tblengine.rows[0].cells[0].paragraphs[0]
        para2=self.table_Tblengine.rows[0].cells[1].paragraphs[0]
        run=para.add_run('엔진명')
        run2=para2.add_run("탐지수")
        run.bold=True
        run2.bold=True
        # self.table_Tblengine.rows[0].cells[0].text = "엔진명"
        # self.table_Tblengine.rows[0].cells[1].text = "탐지수"


    def wrTblthreatpie(self, rownum, cellnum, data):
        self.table_Tblclasspie.rows[0].cells[0].width=Cm(6)
        self.table_Tblclasspie.rows[0].cells[1].width=Cm(3)
        self.table_Tblclasspie.rows[rownum].cells[cellnum].text = data
        
    def wrTblenginechart(self, rownum, cellnum, data):
        self.table_Tblengine.rows[0].cells[0].width=Cm(6)
        self.table_Tblengine.rows[0].cells[1].width=Cm(3)
        self.table_Tblengine.rows[rownum].cells[cellnum].text = data

    def writeDown(self, datas):
        self.doc.add_paragraph(datas)

    def writeTitle(self, datas, level):
        self.doc.add_heading(datas, level=level)

    def writeDownTable(self, datas, num):
        # print(datas[4])
        try:
            num += 1
            self.table.rows[num].cells[0].text = datas[4]
            self.table.rows[num].cells[1].text = datas[10]
            self.table.rows[num].cells[2].text = datas[14]
            self.table.rows[num].cells[3].text = datas[23]
            self.table.rows[num].cells[4].text = datas[16]
            self.table.rows[num].cells[5].text = datas[0] + \
                ">"+datas[1]+">"+datas[2]
        except:
            if datas[0]==0 or datas[1]==0 or datas[2]==0:
                # print(datas)
                pass
            else:
                self.table.rows[num].cells[0].text = datas[4]
                self.table.rows[num].cells[1].text = datas[10]
                self.table.rows[num].cells[2].text = datas[14]
                self.table.rows[num].cells[3].text = datas[23]
                self.table.rows[num].cells[4].text = datas[16]
                self.table.rows[num].cells[5].text = datas[0] + \
                    ">"+datas[1]+">"+datas[2]
            # print(datas)
            # print("err")
    def writeDownTop20Table(self, endpoint, cnt, num):
        num += 1
        self.top20table.rows[num].cells[0].width=Cm(1.5)
        self.top20table.rows[num].cells[2].width=Cm(2.5)
        self.top20table.rows[num].cells[0].text = str(num)
        self.top20table.rows[num].cells[1].text = str(endpoint)
        self.top20table.rows[num].cells[2].text = str(cnt)

        
    def writeDownTable_endpoints(self,key,val,num,cat):
        if cat=="getbyos":
            tbl=self.table_getbyos
        if cat=="getbyver":
            tbl=self.table_getbyver
        if cat=="getbyinf":
            tbl=self.table_getbyinf
        num += 1
        tbl.rows[num].cells[0].width=Cm(4)
        tbl.rows[num].cells[1].width=Cm(4)
        tbl.rows[num].cells[0].text = str(key)
        tbl.rows[num].cells[1].text = str(val)
        
    def writeDownTable_getSolved(self,key,val,num):
        tbl=self.table_getSolved
        num += 1
        tbl.rows[num].cells[0].width=Cm(4)
        tbl.rows[num].cells[1].width=Cm(4)
        tbl.rows[num].cells[0].text = str(key)
        tbl.rows[num].cells[1].text = str(val)
    
    def writeDownTable_thDetail(self,cont,num,num2):
        tbl=self.table_thDetail
        # if num2==0:
        #     tbl.rows[num].cells[num2].width=Cm(4)
        # else:
        #     tbl.rows[num].cells[num2].width=Cm(7)
        tbl.rows[num].cells[num2].text=str(cont)
    
    def writeHyperlink_thDetail(self,name,link):
        tbl=self.table_thDetail
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    print(para.text,"\n")
    
    def addPic(self, picpath):
        self.doc.add_picture(picpath,width=Inches(5),height=Inches(4))
    def addPic2(self, picpath):
        self.doc.add_picture(picpath,width=Inches(4),height=Inches(4))
    def saveDoc(self):
        now = datetime.datetime.now()
        nowDate = now.strftime('%Y%m%d%H%M')
        filepath=self.rootpath+"\\S1reports"
        # directory = os.path(filepath)
        # now = datetime.datetime.now()
        # nowDate = now.strftime('%Y%m%d%H%M%S')
        # nowDate = now.strftime('%Y%m%d')
        try:
            if not os.path.exists(filepath):
                os.makedirs(filepath)
            else:
                scopename=self.scopename.replace('>','_')
                reportdir = filepath+'\\S1Report_'+scopename+"_"+str(nowDate)+'.docx'
                self.doc.save(reportdir)
                self.delPng('./saveBarChart.png')
                self.delPng('./savePieChart.png')
                return reportdir
        except OSError:
            print('Error: Creating directory. ' + filepath)

    def delPng(self,name):
        if os.path.exists(name):
            os.remove(name)