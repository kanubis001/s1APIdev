# 가장 기본적인 기능(문서 열기, 저장, 글자 쓰기 등등)
from docx import Document
from docx.shared import Pt, Cm, Inches
import os
import datetime
now = datetime.datetime.now()
nowDate = now.strftime('%Y%m%d%H%M%S')


class mkReport:
    def __init__(self,startday,endday):
        self.doc = Document()
        style = self.doc.styles['Normal']
        font = style.font
        font.name = '맑은 고딕'
        font.size = Pt(10)
        title='센티넬원 리포트'
        title2='기간 : ',startday,' ~ ',endday
        self.doc.add_heading(title, level=0)
        self.doc.add_heading(title2, level=1)
        self.startday=startday
        self.endday=endday

    def mkTable_detail(self, row, col):
        self.table = self.doc.add_table(rows=row, cols=col)
        self.table.style = self.doc.styles['Table Grid']
        self.table.rows[0].cells[0].text = "agent 명"
        self.table.rows[0].cells[1].text = "IP 주소"
        self.table.rows[0].cells[2].text = "종류"
        self.table.rows[0].cells[3].text = "처리 여부"
        self.table.rows[0].cells[4].text = "위협 구분"
        self.table.rows[0].cells[5].text = "account>site>group"
 
    def mkTop20Table_detail(self, row, col):
        self.top20table = self.doc.add_table(rows=row, cols=col)
        self.top20table.style = self.doc.styles['Table Grid']
        self.top20table.rows[0].cells[0].text = "Endpoint"
        self.top20table.rows[0].cells[1].text = "탐지건수"
       
    def mktable_getbyos(self, row, col):
        self.table_getbyos = self.doc.add_table(rows=row, cols=col)
        self.table_getbyos.style = self.doc.styles['Table Grid']
        self.table_getbyos.rows[0].cells[0].text = "OS 유형"
        self.table_getbyos.rows[0].cells[1].text = "엔드포인트 수량"   
 
    def mktable_getbyver(self, row, col):
        self.table_getbyver = self.doc.add_table(rows=row, cols=col)
        self.table_getbyver.style = self.doc.styles['Table Grid']
        self.table_getbyver.rows[0].cells[0].text = "에이전트 버전"
        self.table_getbyver.rows[0].cells[1].text = "엔드포인트 수량"         

    def mktable_getbyinf(self, row, col):
        self.table_getbyinf = self.doc.add_table(rows=row, cols=col)
        self.table_getbyinf.style = self.doc.styles['Table Grid']
        self.table_getbyinf.rows[0].cells[0].text = "감염여부"
        self.table_getbyinf.rows[0].cells[1].text = "엔드포인트 수량"    
             
    def mkTblthreatpie(self, row, col):
        self.threatTbl = self.doc.add_table(rows=row+1, cols=col)
        self.threatTbl.style = self.doc.styles['Table Grid']
        self.threatTbl.rows[0].cells[0].text = "탐지내역"
        self.threatTbl.rows[0].cells[1].text = "탐지수"
        
    def mkTblenginechart(self, row, col):
        self.threatTbl = self.doc.add_table(rows=row+1, cols=col)
        self.threatTbl.style = self.doc.styles['Table Grid']
        self.threatTbl.rows[0].cells[0].text = "엔진명"
        self.threatTbl.rows[0].cells[1].text = "탐지수" 

    def wrTblthreatpie(self, rownum, cellnum, data):
        self.threatTbl.rows[rownum].cells[cellnum].text = data
        
    def wrTblenginechart(self, rownum, cellnum, data):
        self.threatTbl.rows[rownum].cells[cellnum].text = data

    def writeDown(self, datas):
        self.doc.add_paragraph(datas)

    def writeTitle(self, datas, level):
        self.doc.add_heading(datas, level=level)

    def writeDownTable(self, datas, num):
        # print(datas[4])
        try:
            num += 1
            self.table.rows[num].cells[0].text = datas[4]
            self.table.rows[num].cells[1].text = datas[10]
            self.table.rows[num].cells[2].text = datas[14]
            self.table.rows[num].cells[3].text = datas[23]
            self.table.rows[num].cells[4].text = datas[16]
            self.table.rows[num].cells[5].text = datas[0] + \
                ">"+datas[1]+">"+datas[2]
        except:
            if datas[0]==0 or datas[1]==0 or datas[2]==0:
                # print(datas)
                pass
            else:
                self.table.rows[num].cells[0].text = datas[4]
                self.table.rows[num].cells[1].text = datas[10]
                self.table.rows[num].cells[2].text = datas[14]
                self.table.rows[num].cells[3].text = datas[23]
                self.table.rows[num].cells[4].text = datas[16]
                self.table.rows[num].cells[5].text = datas[0] + \
                    ">"+datas[1]+">"+datas[2]
            # print(datas)
            # print("err")
    def writeDownTop20Table(self, endpoint, cnt, num):
        num += 1
        self.top20table.rows[num].cells[0].text = str(endpoint)
        self.top20table.rows[num].cells[1].text = str(cnt)
    
    def writeDownTable_endpoints(self,key,val,num,cat):
        if cat=="getbyos":
            tbl=self.table_getbyos
        if cat=="getbyver":
            tbl=self.table_getbyver
        if cat=="getbyinf":
            tbl=self.table_getbyinf
        num += 1
        tbl.rows[num].cells[0].text = str(key)
        tbl.rows[num].cells[1].text = str(val)
     
        
    def addPic(self, picpath):
        self.doc.add_picture(picpath,width=Inches(4),height=Inches(4))

    def saveDoc(self):
        # self.doc.save('D:\\원드라이브\\OneDrive\\Develop_private\\Python\\s1ReportingTool\\reporter\\testReport.docx')
        directory = os.path.expanduser('~\\s1report')
        try:
            if not os.path.exists(directory):
                os.makedirs(directory)
            else:

                reportdir = directory+'\\s1report_'+str(nowDate)+'.docx'
                self.doc.save(reportdir)
                return reportdir
        except OSError:
            print('Error: Creating directory. ' + directory)
